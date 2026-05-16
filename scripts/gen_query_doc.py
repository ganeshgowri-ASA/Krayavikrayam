"""Generate the vendor clarification matrix as a .docx.

Produces ``deliverables/<package>-vendor-clarification.docx`` with one
table per vendor showing:

  * Criterion / question
  * Vendor declaration (NC = Not Complied, PC = Partial, CTQ-FC = Full
    Compliance against a CTQ, FC = Full Compliance)
  * Evidence (Ref: file!sheet!Row N) -- copied verbatim from the
    Pinecone retrieval; if the evidence is missing, the cell carries
    "Not declared" and the row is flagged.
"""

from __future__ import annotations

import argparse
import json
import logging
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

from .packages import NOT_DECLARED, by_key
from .scoring import make_ref


log = logging.getLogger("gen_query_doc")

DELIVERABLES_ROOT = Path("deliverables")
EVIDENCE_ROOT = Path("evidence")


CLARIFICATION_ROWS = (
    # (criterion label, evidence field key, expected schema field, classifier)
    ("Compliance tally (FC/PC/DNC)",      "compliance_tally", "fc",            "tally"),
    ("Thermal design",                    "thermal_design",   "refrigerant",   "declared"),
    ("Reliability (MTBF/MTTR/uptime)",    "reliability",      "mtbf",          "declared"),
    ("BoM origin & India stock",          "bom_origin",       "spares",        "bom"),
    ("India service network",             "service_network",  "centres",       "declared"),
    ("Industrial connectivity",           "connectivity",     "opc_ua",        "flags"),
    ("OEM / installed base",              "installed_base",   "oem",           "declared"),
    ("CTQ deviations",                    "ctq_deviations",   "deviations",    "deviations"),
    ("Spectral match (SSS-only)",         "spectral_match",   "spectral_class","declared"),
    ("Lamp / LED life (SSS-only)",        "lamp_life",        "lamp_life_hours","declared"),
)


def _classify(kind: str, field: dict | None) -> str:
    if not field:
        return "NC"
    value = field.get("value", {})
    if kind == "tally":
        try:
            fc = int(value.get("fc", 0) or 0)
            pc = int(value.get("pc", 0) or 0)
            dnc = int(value.get("dnc", 0) or 0)
        except (TypeError, ValueError):
            return "NC"
        if dnc > 0:
            return "NC"
        if pc > 0:
            return "PC"
        if fc > 0:
            return "FC"
        return "NC"
    if kind == "deviations":
        devs = value.get("deviations", [])
        if isinstance(devs, list) and devs:
            return "CTQ-FC" if all(d.get("delta") in (None, 0, "0", "Not declared") for d in devs if isinstance(d, dict)) else "PC"
        return "FC"
    if kind == "flags":
        flags = [k for k in ("opc_ua", "modbus_rtu", "modbus_tcp", "rest_api", "lims") if value.get(k) is True]
        if len(flags) >= 3:
            return "FC"
        if flags:
            return "PC"
        return "NC"
    if kind == "bom":
        spares = value.get("spares", [])
        if not isinstance(spares, list) or not spares:
            return "NC"
        bad = [s for s in spares if isinstance(s, dict) and str(s.get("origin", "")).lower() == "china"
               and (s.get("india_stock_qty") in (None, 0, "", NOT_DECLARED)
                    or s.get("mttr_hours") in (None, 0, "", NOT_DECLARED))]
        return "PC" if bad else "FC"
    # default 'declared'
    if value and any(v not in (None, "", NOT_DECLARED) for v in value.values() if not isinstance(v, (dict, list))):
        return "FC"
    return "NC"


def _render_value(field: dict | None) -> str:
    if not field:
        return NOT_DECLARED
    value = field.get("value", {})
    if not isinstance(value, dict):
        return str(value)
    # Pick the first 3-4 informative keys and render compactly.
    parts = []
    for k, v in value.items():
        if k in ("source_file", "sheet", "row", "row_range", "verbatim_quote"):
            continue
        if isinstance(v, (dict, list)):
            v = json.dumps(v, default=str)[:120]
        parts.append(f"{k}={v}")
        if len(parts) >= 4:
            break
    return "; ".join(parts) if parts else NOT_DECLARED


def generate_doc(
    package_key: str,
    vendors: list[str],
    evidence_root: Path = EVIDENCE_ROOT,
    out_root: Path = DELIVERABLES_ROOT,
) -> Path:
    package = by_key(package_key)
    doc = Document()

    title = doc.add_heading(f"Vendor Clarification Matrix -- {package.display_name}", level=1)
    title.alignment = WD_TABLE_ALIGNMENT.LEFT

    doc.add_paragraph(
        f"Package key: {package.key}. "
        f"Weights: Design Concept = {package.weights.design_concept}, "
        f"Ease of Operation = {package.weights.ease_of_operation}."
    )
    doc.add_paragraph(
        "Compliance codes: FC = Full Compliance, PC = Partial Compliance, "
        "NC = Not Complied, CTQ-FC = Full Compliance against Critical-To-Quality."
    )

    for vendor in vendors:
        doc.add_heading(vendor, level=2)
        ev_path = evidence_root / package.key / f"{_safe(vendor)}.json"
        evidence = json.loads(ev_path.read_text()) if ev_path.exists() else {}

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Criterion"
        hdr[1].text = "Code"
        hdr[2].text = "Vendor declaration"
        hdr[3].text = "Evidence (Ref)"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for label, ev_key, _schema_key, kind in CLARIFICATION_ROWS:
            field = evidence.get(ev_key)
            row = table.add_row().cells
            row[0].text = label
            code = _classify(kind, field)
            row[1].text = code
            row[2].text = _render_value(field)
            row[3].text = make_ref(field)
            if code == "NC":
                for p in row[1].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        r.bold = True
            for cell in row:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

        # Provenance line
        meta = evidence.get("_meta", {})
        para = doc.add_paragraph()
        para.add_run(f"Evidence assembled at {meta.get('generated_at', 'unknown')} "
                     f"from assistant '{meta.get('assistant', 'unknown')}'.").italic = True

    out_root.mkdir(parents=True, exist_ok=True)
    out_path = out_root / f"{package.key}-vendor-clarification.docx"
    doc.save(out_path)
    log.info("Wrote vendor clarification: %s", out_path)
    return out_path


def _safe(name: str) -> str:
    return "".join(c if c.isalnum() or c in "-_." else "_" for c in name)


def main(argv: list[str] | None = None) -> int:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s: %(message)s")
    parser = argparse.ArgumentParser(description="Generate vendor clarification .docx")
    parser.add_argument("--package", required=True)
    parser.add_argument("--vendors", required=True)
    parser.add_argument("--evidence-root", default=str(EVIDENCE_ROOT))
    parser.add_argument("--out-root", default=str(DELIVERABLES_ROOT))
    args = parser.parse_args(argv)
    vendors = [v.strip() for v in args.vendors.split(",") if v.strip()]
    path = generate_doc(
        package_key=args.package,
        vendors=vendors,
        evidence_root=Path(args.evidence_root),
        out_root=Path(args.out_root),
    )
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
